"""Generate a professional .docx proposal document from structured proposal data."""

from __future__ import annotations

import logging
from pathlib import Path

logger = logging.getLogger(__name__)


async def generate_proposal_docx(
    proposal: dict,
    outputs_dir: Path,
    company_name: str,
    diagram_path: Path | None = None,
) -> Path | None:
    """Generate a Word document from the proposal data.

    Args:
        proposal: Structured proposal dict from the Architect Agent.
        outputs_dir: Directory to save the output file.
        company_name: Customer company name (used in filename and cover).
        diagram_path: Optional path to architecture diagram image.

    Returns:
        Path to the generated .docx file, or None if generation failed.
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor
    except ImportError:
        logger.error("python-docx is not installed. Run: uv add python-docx")
        return None

    safe_name = "".join(c if c.isalnum() or c in " _-" else "_" for c in company_name)[:40].strip()
    filename = f"Proposal_{safe_name}.docx" if safe_name else "Proposal.docx"
    output_path = outputs_dir / filename

    doc = Document()

    # ── Styles ───────────────────────────────────────────────────────────────
    BRAND_BLUE = RGBColor(0x16, 0x5B, 0xDA)

    def _set_heading_color(paragraph, color: RGBColor) -> None:
        for run in paragraph.runs:
            run.font.color.rgb = color

    # ── Cover page ───────────────────────────────────────────────────────────
    cover = doc.add_heading("Solution Proposal", level=0)
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_heading_color(cover, BRAND_BLUE)

    subtitle = doc.add_paragraph(f"Khách hàng: {company_name}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)

    powered_by = doc.add_paragraph("Powered by KHantix Multi-Agent AI System")
    powered_by.alignment = WD_ALIGN_PARAGRAPH.CENTER
    powered_by.runs[0].font.size = Pt(10)
    powered_by.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

    # ── Table of Contents placeholder ────────────────────────────────────────
    toc_heading = doc.add_heading("Mục lục", level=1)
    _set_heading_color(toc_heading, BRAND_BLUE)
    toc_items = [
        "1. Executive Summary",
        "2. Phân tích Hiện trạng",
        "3. Kiến trúc Giải pháp",
        "4. Lộ trình Triển khai",
        "5. Ước tính Chi phí & ROI",
        "6. Rủi ro & Mitigation",
    ]
    for item in toc_items:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_page_break()

    # ── 1. Executive Summary ─────────────────────────────────────────────────
    h = doc.add_heading("1. Executive Summary", level=1)
    _set_heading_color(h, BRAND_BLUE)
    doc.add_paragraph(proposal.get("executive_summary", "N/A"))

    # ── 2. Phân tích Hiện trạng ──────────────────────────────────────────────
    h = doc.add_heading("2. Phân tích Hiện trạng", level=1)
    _set_heading_color(h, BRAND_BLUE)
    doc.add_paragraph(proposal.get("current_state_analysis", "N/A"))

    # ── 3. Kiến trúc Giải pháp ───────────────────────────────────────────────
    h = doc.add_heading("3. Kiến trúc Giải pháp", level=1)
    _set_heading_color(h, BRAND_BLUE)

    # Diagram image (if rendered)
    if diagram_path and diagram_path.is_file():
        try:
            doc.add_picture(str(diagram_path), width=Inches(6))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = doc.add_paragraph("Hình 1: Kiến trúc Giải pháp Đề xuất")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.runs[0].font.size = Pt(9)
        except Exception:
            logger.warning("Could not insert diagram image")

    # Mermaid code as fallback/supplement
    arch = proposal.get("solution_architecture", {})
    mermaid_code = arch.get("mermaid_diagram", "")
    if mermaid_code:
        doc.add_heading("Architecture Diagram (Mermaid)", level=2)
        code_para = doc.add_paragraph(mermaid_code)
        code_para.runs[0].font.name = "Courier New"
        code_para.runs[0].font.size = Pt(8)

    # Technology Stack table
    tech_stack = arch.get("tech_stack", [])
    if tech_stack:
        doc.add_heading("Technology Stack", level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Component"
        hdr_cells[1].text = "Technology"
        hdr_cells[2].text = "Mục đích"
        for item in tech_stack:
            row_cells = table.add_row().cells
            row_cells[0].text = str(item.get("component", ""))
            row_cells[1].text = str(item.get("technology", ""))
            row_cells[2].text = str(item.get("purpose", ""))

    # Integration points
    integration_points = arch.get("integration_points", [])
    if integration_points:
        doc.add_heading("Điểm Tích hợp", level=2)
        for pt in integration_points:
            doc.add_paragraph(str(pt), style="List Bullet")

    # ── 4. Lộ trình Triển khai ───────────────────────────────────────────────
    h = doc.add_heading("4. Lộ trình Triển khai", level=1)
    _set_heading_color(h, BRAND_BLUE)
    roadmap = proposal.get("implementation_roadmap", [])
    for phase_info in roadmap:
        phase_num = phase_info.get("phase", "")
        phase_name = phase_info.get("name", "")
        duration = phase_info.get("duration", "")
        objectives = phase_info.get("objectives", "")

        ph = doc.add_heading(f"Phase {phase_num}: {phase_name}", level=2)
        if duration:
            doc.add_paragraph(f"Thời gian: {duration}")
        if objectives:
            doc.add_paragraph(f"Mục tiêu: {objectives}")

        tasks = phase_info.get("tasks", [])
        if tasks:
            doc.add_paragraph("Công việc chính:")
            for task in tasks:
                doc.add_paragraph(str(task), style="List Bullet")

        deliverables = phase_info.get("deliverables", [])
        if deliverables:
            doc.add_paragraph("Deliverables:")
            for d in deliverables:
                doc.add_paragraph(str(d), style="List Bullet")

    # ── 5. Ước tính Chi phí & ROI ─────────────────────────────────────────────
    h = doc.add_heading("5. Ước tính Chi phí & ROI", level=1)
    _set_heading_color(h, BRAND_BLUE)
    cost = proposal.get("cost_estimation", {})

    phases_cost = cost.get("phases", [])
    if phases_cost:
        cost_table = doc.add_table(rows=1, cols=4)
        cost_table.style = "Light Grid Accent 1"
        hdr = cost_table.rows[0].cells
        hdr[0].text = "Phase"
        hdr[1].text = "Setup ($)"
        hdr[2].text = "Monthly OpEx ($)"
        hdr[3].text = "Implementation ($)"
        for p in phases_cost:
            r = cost_table.add_row().cells
            r[0].text = f"Phase {p.get('phase', '')}: {p.get('name', '')}"
            r[1].text = f"{p.get('setup_usd', 0):,}" if p.get("setup_usd") else "-"
            r[2].text = f"{p.get('monthly_usd', 0):,}" if p.get("monthly_usd") else "-"
            r[3].text = f"{p.get('implementation_usd', 0):,}" if p.get("implementation_usd") else "-"

    totals = []
    if cost.get("total_implementation_usd"):
        totals.append(f"Tổng chi phí Implementation: ${cost['total_implementation_usd']:,}")
    if cost.get("total_monthly_opex_usd"):
        totals.append(f"Monthly OpEx (sau khi go-live): ${cost['total_monthly_opex_usd']:,}/tháng")
    if cost.get("roi_payback_months"):
        totals.append(f"Ước tính payback period: {cost['roi_payback_months']} tháng")
    for t in totals:
        doc.add_paragraph(t)

    assumptions = cost.get("assumptions", [])
    if assumptions:
        doc.add_heading("Assumptions", level=2)
        for a in assumptions:
            doc.add_paragraph(str(a), style="List Bullet")

    # ── 6. Rủi ro & Mitigation ───────────────────────────────────────────────
    h = doc.add_heading("6. Rủi ro & Mitigation", level=1)
    _set_heading_color(h, BRAND_BLUE)
    risks = proposal.get("risks", [])
    if risks:
        risk_table = doc.add_table(rows=1, cols=4)
        risk_table.style = "Light Grid Accent 1"
        hdr = risk_table.rows[0].cells
        hdr[0].text = "Rủi ro"
        hdr[1].text = "Mức độ"
        hdr[2].text = "Xác suất"
        hdr[3].text = "Kế hoạch Giảm thiểu"
        for r in risks:
            row = risk_table.add_row().cells
            row[0].text = str(r.get("risk", ""))
            row[1].text = str(r.get("level", ""))
            row[2].text = str(r.get("probability", ""))
            row[3].text = str(r.get("mitigation", ""))

    # ── Footer ───────────────────────────────────────────────────────────────
    doc.add_page_break()
    footer_para = doc.add_paragraph("Document generated by KHantix Multi-Agent AI System (DeerFlow POC)")
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.runs[0].font.size = Pt(9)
    footer_para.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.save(str(output_path))
    logger.info("[DocxGenerator] Proposal saved to %s", output_path)
    return output_path
